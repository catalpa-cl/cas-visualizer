from __future__ import annotations

import abc
from bisect import bisect_left, bisect_right
from dataclasses import dataclass, field
from itertools import cycle
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional

import pandas as pd
from cassis import Cas, TypeSystem
from cassis.typesystem import FeatureStructure
from spacy.displacy import EntityRenderer, DependencyRenderer, SpanRenderer

from cas_visualizer.util import ensure_typesystem


class VisualizerException(Exception):
    """
    Domain-specific error for visualizer configuration and rendering.

    Raised when:
    - configuration is invalid (unknown type/feature, illegal span type),
    - no content can be rendered in strict mode,
    - requested output format is unsupported,
    - spec cannot be produced (e.g., misaligned annotation boundaries).
    """
    pass


# ---------- Type configuration ----------

@dataclass
class TypeConfig:
    """
    Per-type configuration used by visualizers.

    Fields:
    - name: fully qualified CAS type name (e.g., "de.tudarmstadt.ukp.dkpro.core.api.ner.type.NamedEntity")
    - feature: feature name to use for labels (e.g., "value"); if None, no label
    - default_color: fallback color for this type
    - default_label: fallback label if feature value is missing
    - value_labels: raw feature value -> display label mapping
    - label_colors: display label -> color mapping
    """
    name: str
    feature: str | None = None
    default_color: str | None = None
    default_label: str | None = None
    value_labels: Dict[Any, str] = field(default_factory=dict)
    label_colors: Dict[str, str] = field(default_factory=dict)

    def label_for_value(self, value: Any) -> str | None:
        """
        Resolve a display label:
        - use value_labels mapping if present,
        - else str(value) if value is not None,
        - else fallback to default_label.
        """
        mapped = self.value_labels.get(value)
        if mapped is not None:
            return mapped
        if value is not None:
            return str(value)
        return self.default_label

    def color_for_label(self, label: str | None) -> str | None:
        """
        Resolve a color:
        - use label_colors[label] if present,
        - else fallback to default_color.
        """
        if label is None:
            return self.default_color
        return self.label_colors.get(label, self.default_color)


# ---------- Base visualizer ----------

class Visualizer(abc.ABC):
    """
    Base class for CAS visualizers.

    Responsibilities:
    - hold per-type configuration (features, labels, colors),
    - provide helpers to resolve labels and colors,
    - normalize the type system argument via ensure_typesystem,
    - offer an HTML page wrapper (UTF‑8 meta) for fragments.

    Contract:
    - visualize and render return a single string, always.
    - Subclasses validate supported output_format values at runtime and raise VisualizerException otherwise.
    """
    def __init__(self, ts: str | Path | TypeSystem):
        # Registered type configs keyed by full type path
        self._type_configs: Dict[str, TypeConfig] = {}

        # Default color iterator (cycled, never exhausted)
        color_list = [
            "lightgreen", "orangered", "orange", "plum", "palegreen",
            "mediumseagreen", "steelblue", "skyblue", "navajowhite",
            "mediumpurple", "rosybrown", "silver", "gray", "paleturquoise",
        ]
        self._default_colors: Iterator[str] = cycle(color_list)

        # Normalize and store the type system
        self._ts = ensure_typesystem(ts)

    @property
    def ts(self) -> TypeSystem:
        """The normalized TypeSystem used by this visualizer."""
        return self._ts

    # ----- Configuration API -----

    def add_type(
        self,
        name: str,
        feature: str | None = None,
        color: str | None = None,
        label: str | None = None,
    ) -> None:
        """
        Register or update a type configuration.

        Parameters:
        - name: fully qualified CAS type name
        - feature: feature to use for labels (optional)
        - color: default color (optional; if omitted, a new default is assigned)
        - label: default label (optional; if omitted, last path segment is used)
        """
        if not name:
            raise VisualizerException("type path cannot be empty")

        cfg = self._type_configs.get(name)
        if cfg is None:
            cfg = TypeConfig(name=name)
            self._type_configs[name] = cfg

        if feature is not None:
            # Optional: validate against TypeSystem (expensive if repeated)
            # ts_type = self._ts.get_type(name)
            # if feature not in {f.name for f in ts_type.features}:
            #     raise VisualizerException(f'Unknown feature "{feature}" for type "{name}"')
            cfg.feature = feature

        if color is not None:
            cfg.default_color = color
        elif cfg.default_color is None:
            cfg.default_color = next(self._default_colors)

        cfg.default_label = label if label is not None else (cfg.default_label or name.split(".")[-1])

    def add_feature(
        self,
        name: str,
        feature: str,
        value: Any,
        color: str | None = None,
        label: str | None = None,
    ) -> None:
        """
        Map a specific feature value to a display label and color.

        Parameters:
        - name: CAS type name to configure
        - feature: feature on the type (stored on the type config)
        - value: raw feature value to map
        - color: optional color for the display label (assigned if given; otherwise, a default is picked)
        - label: optional display label (defaults to str(value))
        """
        if not name:
            raise VisualizerException("type name cannot be empty")

        cfg = self._type_configs.get(name)
        if cfg is None:
            cfg = TypeConfig(name=name)
            self._type_configs[name] = cfg

        if not feature:
            raise VisualizerException(f"a feature for type {name} must be specified")
        cfg.feature = feature

        if value is None:
            raise VisualizerException(f"a value for feature {feature} must be specified")

        eff_label = label if label is not None else str(value)
        cfg.value_labels[value] = eff_label

        if color is not None:
            cfg.label_colors[eff_label] = color
        elif eff_label not in cfg.label_colors:
            cfg.label_colors[eff_label] = next(self._default_colors)

        if cfg.default_color is None:
            cfg.default_color = next(self._default_colors)

    def remove_type(self, type_path: str) -> None:
        """Remove a type configuration and all associated mappings."""
        if not type_path:
            raise VisualizerException("type path cannot be empty")
        self._type_configs.pop(type_path, None)

    def clear_types(self) -> None:
        """Remove all type configurations."""
        self._type_configs.clear()

    # ----- Resolvers / Helpers -----

    def list_types(self) -> list[str]:
        """List configured type names (in insertion order)."""
        return list(self._type_configs.keys())

    def iter_type_configs(self):
        """Iterate over (type_name, TypeConfig) pairs."""
        return self._type_configs.items()

    def get_type_feature(self, type_name: str) -> str | None:
        """Return the configured feature name for a type (or None)."""
        cfg = self._type_configs.get(type_name)
        return cfg.feature if cfg else None

    def resolve_label(self, fs: FeatureStructure, type_name: str) -> str | None:
        """Compute a display label for a feature structure of the given type."""
        cfg = self._type_configs.get(type_name)
        if cfg is None:
            return None
        value = self._get_feature_value(fs, cfg.feature)
        return cfg.label_for_value(value)

    def resolve_color(self, type_name: str, label: str | None) -> str | None:
        """Compute a display color for the given type and label."""
        cfg = self._type_configs.get(type_name)
        if cfg is None:
            return None
        return cfg.color_for_label(label)

    def _get_feature_value(self, fs: FeatureStructure, feature_name: str | None) -> Any:
        """Safely get a feature value (raises VisualizerException on missing feature)."""
        if feature_name is None:
            return None
        try:
            return fs.get(feature_name)
        except KeyError as e:
            raise VisualizerException(
                f"Feature '{feature_name}' not found on type '{fs.type.name}'"
            ) from e

    @staticmethod
    def _wrap_html_page(fragment: str, title: str = "Visualizer") -> str:
        """
        Wrap an HTML fragment into a full document with UTF‑8 meta.

        Useful to avoid browser encoding issues for fragments rendered by spaCy's displaCy.
        """
        return (
            "<!doctype html>\n"
            "<html lang=\"en\">\n"
            "<head>\n<meta charset=\"utf-8\">\n"
            f"<title>{title}</title>\n</head>\n<body>\n"
            f"{fragment}\n"
            "</body>\n</html>"
        )

    @abc.abstractmethod
    def visualize(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
        output_format: str = "html",
    ) -> str:
        """
        Build and render a visualization.

        Parameters:
        - cas: CAS instance to visualize
        - start, end: character range to consider (end = -1 means end of document)
        - output_format: format identifier (subclasses validate supported values)

        Returns:
        - a single string (HTML, SVG, CSV, JSON, LaTeX depending on the visualizer and selected format)

        Errors:
        - VisualizerException on unsupported format or when strict mode requires content but none is found.
        """
        raise NotImplementedError


# ---------- Table visualizer ----------

class TableVisualizer(Visualizer):
    """
    Tabular visualization of selected annotations.

    Pipeline:
    - build: produce a pandas DataFrame of rows within [start, end)
    - render: export the DataFrame as HTML, CSV, JSON, or LaTeX (returns str)
    - visualize: build + render

    Strict mode:
    - If strict=True and the DataFrame is empty, render() raises VisualizerException.
      Otherwise, you may get a valid but empty HTML table/CSV/JSON/LaTeX.
    """
    def __init__(
        self,
        ts: str | Path | TypeSystem,
        *,
        default_render_options: Dict[str, Any] | None = None,
        default_sort: bool = True,
        page: bool = False,
        strict: bool = True,
    ):
        super().__init__(ts)
        self._default_render_options = default_render_options or {}
        self._default_sort = default_sort
        self._page = page
        self._strict = strict

    def build(self, cas: Cas, *, start: int = 0, end: int = -1) -> pd.DataFrame:
        """
        Build the table spec (DataFrame). Include only annotations fully inside [start, end).
        If end == -1, treat it as len(document_text).
        """
        if end == -1:
            end = len(cas.sofa_string)
        cols = ["text", "feature", "value", "begin", "end"]
        records: list[dict[str, Any]] = []

        for type_name, cfg in self.iter_type_configs():
            for fs in cas.select(type_name):
                if fs.begin < start or fs.end > end:
                    continue

                feature_name = cfg.feature
                feature_value = self._get_feature_value(fs, feature_name)

                records.append({
                    "text": fs.get_covered_text(),
                    "feature": feature_name,
                    "value": feature_value,
                    "begin": fs.begin,
                    "end": fs.end,
                })

        df = pd.DataFrame.from_records(records, columns=cols)
        if self._default_sort and not df.empty:
            df = df.sort_values(by=["begin", "end"], kind="mergesort")
        return df

    def render(
        self,
        spec: pd.DataFrame,
        *,
        output_format: str = "html",
        render_options: Dict[str, Any] | None = None,
    ) -> str:
        """
        Export the table to the requested format.

        Supported formats:
        - 'html': returns an HTML fragment; wrap to a full page if page=True
        - 'csv': returns CSV text (no index)
        - 'json': returns JSON (records orientation)
        - 'latex': returns LaTeX tabular code

        Errors:
        - VisualizerException on unsupported format or empty result in strict mode.
        """
        fmt = output_format.lower()
        opts = {**self._default_render_options, **(render_options or {})}

        if spec.empty and self._strict:
            raise VisualizerException("TableVisualizer: empty result (no rows in selected range).")

        if fmt == "html":
            frag = spec.to_html(**({"index": False, "escape": True} | opts))
            return self._wrap_html_page(frag, "TableVisualizer") if self._page else frag

        if fmt == "csv":
            return spec.to_csv(**({"index": False} | opts))

        if fmt == "json":
            return spec.to_json(**({"orient": "records"} | opts))

        if fmt == "latex":
            return spec.to_latex(**({"index": False, "escape": True} | opts))

        raise VisualizerException(f"Unsupported table output format: {fmt}")

    def visualize(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
        output_format: str = "html",
    ) -> str:
        """Convenience wrapper: build + render."""
        spec = self.build(cas, start=start, end=end)
        return self.render(spec, output_format=output_format)


# ---------- Span visualizer ----------

class SpacySpanVisualizer(Visualizer):
    """
    Span visualization using spaCy’s displaCy.

    Modes:
    - HIGHLIGHT: render entity-like colored spans via EntityRenderer
    - UNDERLINE: render underlined spans via SpanRenderer (token-based)

    Notes:
    - displaCy’s color map is keyed by label strings only, so different types with the same label share colors.
      To avoid collisions, prefer unique labels across types.
    """
    HIGHLIGHT: str = "HIGHLIGHT"
    UNDERLINE: str = "UNDERLINE"

    def __init__(
        self,
        ts: str | Path | TypeSystem,
        span_type: str | None = None,
        types: list[str] | None = None,
        *,
        page: bool = False,
        strict: bool = True,
    ):
        super().__init__(ts)
        self._span_types: list[str] = [SpacySpanVisualizer.HIGHLIGHT, SpacySpanVisualizer.UNDERLINE]
        self._selected_span_type = SpacySpanVisualizer.UNDERLINE
        if span_type is not None:
            self.selected_span_type = span_type
        self._allow_highlight_overlap = False
        self._page = page
        self._strict = strict
        if types is not None:
            for type_name in types:
                self.add_type(type_name)

    @property
    def selected_span_type(self) -> str:
        """Current visualization mode: 'HIGHLIGHT' or 'UNDERLINE'."""
        return self._selected_span_type

    @selected_span_type.setter
    def selected_span_type(self, value: str) -> None:
        if value not in self._span_types:
            raise VisualizerException(f"Invalid span type: {value}. Expected one of {self._span_types}")
        self._selected_span_type = value

    @property
    def allow_highlight_overlap(self) -> bool:
        """Whether overlapping highlights are allowed (HIGHLIGHT mode)."""
        return self._allow_highlight_overlap

    @allow_highlight_overlap.setter
    def allow_highlight_overlap(self, value: bool):
        self._allow_highlight_overlap = value

    # ------------- build / render / visualize -------------

    def build(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
    ) -> dict[str, Any]:
        """
        Build a displaCy spec for the selected span visualization.

        Returns:
        - HIGHLIGHT: {'mode': 'HIGHLIGHT', 'text': str, 'ents': list[dict], 'colors': dict[label, color]}
        - UNDERLINE: {'mode': 'UNDERLINE', 'token_texts': list[str], 'spans': list[dict], 'colors': dict[label, color]}

        Errors:
        - VisualizerException if overlapping highlights are detected (and overlap not allowed) or no content in strict mode.
        """
        if end == -1:
            end = len(cas.sofa_string)

        if self.selected_span_type == SpacySpanVisualizer.HIGHLIGHT:
            tmp_ents: list[dict[str, Any]] = []
            labels_to_colors: dict[str, str] = {}
            for annotation_type in self.list_types():
                for fs in cas.select(annotation_type):
                    if fs.begin < start or fs.end > end:
                        continue
                    label = self.resolve_label(fs, annotation_type)
                    color = self.resolve_color(annotation_type, label)
                    if color:
                        tmp_ents.append({"start": fs.begin, "end": fs.end, "label": label})
                        labels_to_colors[label] = color

            sorted_ents = sorted(tmp_ents, key=lambda x: (x["start"], x["end"]))
            if not self._allow_highlight_overlap and self._check_overlap(sorted_ents):
                raise VisualizerException(
                    "Highlighted annotations are overlapping. Choose a different set of annotations "
                    "or set allow_highlight_overlap=True."
                )
            if self._strict and not sorted_ents:
                raise VisualizerException(
                    f"No entities found for configured types {self.list_types()} in range [{start}, {end}]."
                )

            return {
                "mode": SpacySpanVisualizer.HIGHLIGHT,
                "text": cas.sofa_string,
                "ents": sorted_ents,
                "colors": labels_to_colors,
            }

        elif self.selected_span_type == SpacySpanVisualizer.UNDERLINE:
            selected_annotations = [
                item for typeclass in self.list_types() for item in cas.select(typeclass)
                if item.begin >= start and item.end <= end
            ]
            if self._strict and not selected_annotations:
                raise VisualizerException(
                    f"No spans found for configured types {self.list_types()} in range [{start}, {end}]."
                )

            tmp_tokens = self._split_text_for_spans(cas.sofa_string, selected_annotations)
            tmp_token_texts = [_["text"] for _ in sorted(tmp_tokens, key=lambda t: t["start"])]

            tmp_spans: list[dict[str, Any]] = []
            labels_to_colors: dict[str, str] = {}
            for annotation_type in self.list_types():
                for tmp_span in self._create_spans(
                    cas=cas, cas_sofa_tokens=tmp_tokens, annotation_type=annotation_type, start=start, end=end
                ):
                    label = tmp_span["label"]
                    color = self.resolve_color(annotation_type, label)
                    if color is not None:
                        labels_to_colors[label] = color
                        tmp_spans.append(tmp_span)

            tmp_spans.sort(key=lambda x: x["start"])
            if self._strict and not tmp_spans:
                raise VisualizerException(
                    f"No underline spans could be created for configured types {self.list_types()} in range [{start}, {end}]."
                )

            return {
                "mode": SpacySpanVisualizer.UNDERLINE,
                "token_texts": tmp_token_texts,
                "spans": tmp_spans,
                "colors": labels_to_colors,
            }

        else:
            raise VisualizerException("Invalid span type")

    def render(
        self,
        spec: dict[str, Any],
        *,
        output_format: str = "html",
    ) -> str:
        """
        Render the built spec using spaCy’s displaCy.

        Supported formats:
        - 'html' only (returns a string). Set page=True to get a full HTML document.

        Errors:
        - VisualizerException on unsupported format or invalid spec mode.
        """
        fmt = output_format.lower()
        if fmt != "html":
            raise VisualizerException("SpacySpanVisualizer supports only 'html' output_format")

        mode = spec.get("mode")
        if mode == SpacySpanVisualizer.HIGHLIGHT:
            frag = EntityRenderer({"colors": spec["colors"]}).render_ents(spec["text"], spec["ents"], "")
        elif mode == SpacySpanVisualizer.UNDERLINE:
            frag = SpanRenderer({"colors": spec["colors"]}).render_spans(spec["token_texts"], spec["spans"], "")
        else:
            raise VisualizerException("Invalid spec: missing or unknown mode")

        return self._wrap_html_page(frag, "SpanVisualizer") if self._page else frag

    def visualize(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
        output_format: str = "html",
    ) -> str:
        """Convenience: build + render (returns str)."""
        spec = self.build(cas, start=start, end=end)
        return self.render(spec, output_format=output_format)

    # ------------- private helpers -------------

    @staticmethod
    def _check_overlap(sorted_ents: list[dict[str, Any]]) -> bool:
        """
        Detect overlap in a list of spans sorted by (start, end).
        Returns True if any span starts before the previous one ended.
        """
        prev_end = -1
        for ent in sorted_ents:
            if ent["start"] < prev_end:
                return True
            prev_end = max(prev_end, ent["end"])
        return False

    @staticmethod
    def _split_text_for_spans(
        cas_sofa_string: str,
        feature_structures: list[FeatureStructure],
    ) -> list[dict[str, int | str]]:
        """
        Split text at annotation boundaries and after whitespace to build token-like chunks for SpanRenderer.
        Returns list of dicts: {start, end, text}.
        """
        cutting_points = {fs.begin for fs in feature_structures} | {fs.end for fs in feature_structures}
        cutting_points |= {i + 1 for i, ch in enumerate(cas_sofa_string) if ch.isspace()}

        points = sorted(cutting_points | {0, len(cas_sofa_string)})

        tokens: list[dict[str, int | str]] = []
        for i in range(len(points) - 1):
            start, end = points[i], points[i + 1]
            if start == end:
                continue
            tokens.append({"start": start, "end": end, "text": cas_sofa_string[start:end]})

        return tokens

    def _create_spans(
        self,
        cas: Cas,
        cas_sofa_tokens: list[dict[str, int | str]],
        annotation_type: str,
        *,
        start: int,
        end: int,
    ) -> list[dict[str, int | str | None]]:
        """
        Project annotation character offsets onto token indices for the SpanRenderer.

        Returns dicts with:
        - start/end (char offsets),
        - start_token/end_token (token indices),
        - label (resolved via type configuration).
        """
        start_idx = {tok["start"]: i for i, tok in enumerate(cas_sofa_tokens)}
        end_idx_next = {tok["end"]: i + 1 for i, tok in enumerate(cas_sofa_tokens)}  # exclusive end

        spans: list[dict[str, int | str | None]] = []
        for fs in cas.select(annotation_type):
            if fs.begin < start or fs.end > end:
                continue
            try:
                start_tok = start_idx[fs.begin]
                end_tok = end_idx_next[fs.end]
            except KeyError as e:
                raise VisualizerException(
                    f"Annotation [{fs.begin}, {fs.end}] of type {annotation_type} "
                    f"does not align with token boundaries."
                ) from e

            spans.append(
                {
                    "start": fs.begin,
                    "end": fs.end,
                    "start_token": start_tok,
                    "end_token": end_tok,
                    "label": self.resolve_label(fs, annotation_type),
                }
            )
        return spans


# ---------- Dependency visualizer ----------

class SpacyDependencyVisualizer(Visualizer):
    """
    Dependency visualization inside sentence spans using spaCy’s displaCy.

    Contract:
    - build: produce a list of per-sentence specs [{'words': [...], 'arcs': [...]}]
    - render: 'html' → one HTML string (multiple sentences supported),
              'svg'  → one SVG string (requires exactly one sentence/spec)
    - visualize: build + render (returns str)

    Strict mode:
    - If strict=True and there are no sentences or no tokens/arcs, build() raises VisualizerException.
    """
    def __init__(
        self,
        ts: TypeSystem | Path | str,
        dep_type: str,
        pos_type: str,
        span_type: str,
        *,
        minify: bool = False,
        options: Dict[str, Any] | None = None,
        page: bool = False,
        strict: bool = True,
    ):
        super().__init__(ts)
        self._dep_type = dep_type
        self._pos_type = pos_type
        self._span_type = span_type
        self._minify = minify
        self._options = options or {}
        self._page = page
        self._strict = strict

    def build(self, cas: Cas, *, start: int = 0, end: int = -1) -> list[dict[str, Any]]:
        """
        Build displaCy specs for all sentence spans within [start, end].

        Returns:
        - list of dicts, one per sentence: {'words': [...], 'arcs': [...]}

        Errors:
        - VisualizerException if no sentence spans found or (in strict mode) no tokens/arcs.
        """
        if end == -1:
            end = len(cas.sofa_string)
        if end >= 0 and start > end:
            raise VisualizerException(f"Given span range [start={start}, end={end}] is not valid.")

        parsed: list[dict[str, Any]] = []
        for sent in cas.select(self._span_type):
            if sent.begin >= start and sent.end <= end:
                struct = self._dep_to_dict(cas=cas, covered=sent)
                parsed.append({"words": struct["words"], "arcs": struct["arcs"]})

        if not parsed:
            raise VisualizerException(f"No spans found for type {self._span_type} in range [{start}, {end}].")

        total_words = sum(len(p.get("words", [])) for p in parsed)
        total_arcs = sum(len(p.get("arcs", [])) for p in parsed)
        if self._strict and (total_words == 0 or total_arcs == 0):
            raise VisualizerException(
                f"DependencyVisualizer: nothing to render (words={total_words}, arcs={total_arcs}) "
                f"in range [{start}, {end}]."
            )

        return parsed

    def render(
        self,
        spec: list[dict[str, Any]],
        *,
        output_format: str = "html",
    ) -> str:
        """
        Render the displaCy spec.

        Supported formats:
        - 'html': returns an HTML fragment combining all sentence specs; wrap to full page if page=True.
        - 'svg': returns a single SVG string; exactly one spec required.

        Errors:
        - VisualizerException on unsupported format or if 'svg' is requested for multiple sentences.
        """
        fmt = output_format.lower()
        renderer = DependencyRenderer(options=self._options)

        if fmt == "html":
            frag = renderer.render(spec, page=False, minify=self._minify)
            return self._wrap_html_page(frag, "DependencyVisualizer") if self._page else frag

        if fmt == "svg":
            if len(spec) != 1:
                raise VisualizerException(
                    f"SVG output supports exactly one spec; got {len(spec)}. "
                    f"Render per sentence (build → choose one), or use output_format='html' for multiple sentences."
                )
            p = spec[0]
            return renderer.render_svg("render_id-0", p["words"], p["arcs"])

        raise VisualizerException(f"Output format {output_format} is not supported.")

    def visualize(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
        output_format: str = "html",
    ) -> str:
        """Convenience: build + render (returns str)."""
        spec = self.build(cas, start=start, end=end)
        return self.render(spec, output_format=output_format)

    def _dep_to_dict(self, cas: Cas, covered: FeatureStructure) -> dict[str, Any]:
        """
        Build a displaCy-compatible structure for one sentence.

        Returns:
        - dict with 'words' (tokens with POS) and 'arcs' (dependencies with direction/labels).
        """
        covered_pos = list(cas.select_covered(self._pos_type, covering_annotation=covered))
        offset_to_index = {p.begin: i for i, p in enumerate(covered_pos)}

        words = [
            {"text": p.get_covered_text(), "tag": getattr(p, "PosValue", None)}
            for p in covered_pos
        ]

        cbegin, cend = covered.begin, covered.end
        arcs: list[dict[str, Any]] = []
        for d in cas.select(self._dep_type):
            gb = d.Governor.begin
            db = d.Dependent.begin
            if (cbegin <= gb < cend) and (cbegin <= db < cend):
                if gb in offset_to_index and db in offset_to_index:
                    start_idx = offset_to_index[gb]
                    end_idx = offset_to_index[db]
                    dir_ = "right" if gb <= db else "left"
                    # Normalize to start <= end; keep direction in 'dir'
                    if start_idx > end_idx:
                        start_idx, end_idx = end_idx, start_idx
                    # Avoid self-loops
                    if start_idx != end_idx:
                        arcs.append({
                            "start": start_idx,
                            "end": end_idx,
                            "label": d.DependencyType,
                            "dir": dir_,
                        })

        return {"words": words, "arcs": arcs}

# ---------- DOCX span visualizer (Word document with highlighting/underline) ----------

class DocxSpanVisualizer(Visualizer):
    """
    Span visualization that outputs a Word (.docx) document with text and highlighting/underline.

    Modes:
    - HIGHLIGHT: highlight spans with background colors (limited to Word's WD_COLOR_INDEX palette).
    - UNDERLINE: underline spans and optionally set the text color.

    Contract:
    - build: similar to SpacySpanVisualizer; gathers spans within [start, end].
      Returns a spec dict: {'mode': 'HIGHLIGHT'|'UNDERLINE', 'text': str, 'spans': list[dict]}
        Each span dict: {'start': int, 'end': int, 'label': str|None, 'color': str|None}
    - render: supports 'docx' only; returns base64-encoded DOCX bytes as a string.
    - visualize: build + render (returns the base64-encoded DOCX string).

    Strictness and overlap:
    - If strict=True and no spans are found, a VisualizerException is raised.
    - In HIGHLIGHT mode with allow_highlight_overlap=False, overlapping spans raise VisualizerException.
      If allow_highlight_overlap=True, overlapping regions will be formatted with the most recently begun span.

    Dependencies:
    - Requires `python-docx` (pip install python-docx).
    """

    HIGHLIGHT: str = "HIGHLIGHT"
    UNDERLINE: str = "UNDERLINE"

    def __init__(
        self,
        ts: str | Path | TypeSystem,
        *,
        mode: str = HIGHLIGHT,
        types: list[str] | None = None,
        allow_highlight_overlap: bool = False,
        strict: bool = True,
    ):
        super().__init__(ts)
        if mode not in (DocxSpanVisualizer.HIGHLIGHT, DocxSpanVisualizer.UNDERLINE):
            raise VisualizerException(
                f"Invalid mode: {mode}. Expected one of [{DocxSpanVisualizer.HIGHLIGHT}, {DocxSpanVisualizer.UNDERLINE}]"
            )
        self._mode = mode
        self._allow_highlight_overlap = allow_highlight_overlap
        self._strict = strict
        if types is not None:
            for type_name in types:
                self.add_type(type_name)

    # ------------- build / render / visualize -------------

    def build(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
    ) -> dict[str, Any]:
        """
        Build a spec for DOCX rendering.

        Returns:
        - {'mode': 'HIGHLIGHT'|'UNDERLINE', 'text': str, 'spans': list[dict]}
          where each span dict has: {'start': int, 'end': int, 'label': str|None, 'color': str|None}

        Errors:
        - VisualizerException if overlapping highlights are detected (and overlap not allowed) or no content in strict mode.
        """
        if end == -1:
            end = len(cas.sofa_string)

        spans: list[dict[str, Any]] = []
        for annotation_type in self.list_types():
            for fs in cas.select(annotation_type):
                if fs.begin < start or fs.end > end:
                    continue
                label = self.resolve_label(fs, annotation_type)
                color = self.resolve_color(annotation_type, label)
                spans.append({"start": fs.begin, "end": fs.end, "label": label, "color": color})

        spans.sort(key=lambda s: (s["start"], s["end"]))

        if self._mode == DocxSpanVisualizer.HIGHLIGHT and not self._allow_highlight_overlap:
            if self._check_overlap(spans):
                raise VisualizerException(
                    "Highlighted annotations are overlapping. Choose a different set of annotations "
                    "or enable allow_highlight_overlap=True."
                )

        if self._strict and not spans:
            raise VisualizerException(
                f"No spans found for configured types {self.list_types()} in range [{start}, {end}]."
            )

        return {"mode": self._mode, "text": cas.sofa_string, "spans": spans}

    def render(
        self,
        spec: dict[str, Any],
        *,
        output_format: str = "docx",
    ) -> str:
        """
        Render the built spec to an in-memory .docx and return base64-encoded bytes.

        Supported formats:
        - 'docx' only

        Returns:
        - base64-encoded DOCX (str)

        Errors:
        - VisualizerException on unsupported format, invalid spec mode, or missing python-docx.
        """
        fmt = output_format.lower()
        if fmt != "docx":
            raise VisualizerException("DocxSpanVisualizer supports only 'docx' output_format")

        try:
            from io import BytesIO
            from base64 import b64encode
            from docx import Document  # python-docx
            from docx.enum.text import WD_COLOR_INDEX
            from docx.shared import RGBColor
        except ImportError as e:
            raise VisualizerException("python-docx is required for DOCX rendering (pip install python-docx).") from e

        mode = spec.get("mode")
        if mode not in (DocxSpanVisualizer.HIGHLIGHT, DocxSpanVisualizer.UNDERLINE):
            raise VisualizerException("Invalid spec: missing or unknown mode")

        text: str = spec["text"]
        spans: list[dict[str, Any]] = spec["spans"]

        # Prepare events for a sweep-line over the text to create formatted runs
        events: list[tuple[int, str, dict[str, Any]]] = []  # (pos, 'begin'|'end', span)
        for s in spans:
            events.append((s["start"], "begin", s))
            events.append((s["end"], "end", s))
        # Process end events before begin events at the same position to avoid zero-width segments
        events.sort(key=lambda t: (t[0], 0 if t[1] == "end" else 1))

        doc = Document()
        para = doc.add_paragraph()

        active: list[dict[str, Any]] = []
        prev_pos = 0

        for pos, kind, span in events:
            if pos > prev_pos:
                segment_text = text[prev_pos:pos]
                run = para.add_run(segment_text)
                # Apply style based on current active spans
                style_span = active[-1] if active else None  # last-opened span wins for overlap
                if style_span is not None:
                    color_name = style_span.get("color")
                    if mode == DocxSpanVisualizer.HIGHLIGHT:
                        run.font.highlight_color = self._map_to_highlight(color_name, WD_COLOR_INDEX)  # type: ignore[attr-defined]
                    elif mode == DocxSpanVisualizer.UNDERLINE:
                        run.font.underline = True
                        rgb = self._map_to_rgb(color_name, RGBColor)
                        if rgb is not None:
                            run.font.color.rgb = rgb
                prev_pos = pos

            # Update active set
            if kind == "begin":
                active.append(span)
            else:  # 'end'
                # remove the first matching span by identity or by (start,end,label)
                for i, a in enumerate(active):
                    if a is span or (a["start"], a["end"], a.get("label")) == (span["start"], span["end"], span.get("label")):
                        active.pop(i)
                        break

        # Tail segment
        if prev_pos < len(text):
            run = para.add_run(text[prev_pos:])
            style_span = active[-1] if active else None
            if style_span is not None:
                color_name = style_span.get("color")
                if mode == DocxSpanVisualizer.HIGHLIGHT:
                    run.font.highlight_color = self._map_to_highlight(color_name, WD_COLOR_INDEX)  # type: ignore[attr-defined]
                elif mode == DocxSpanVisualizer.UNDERLINE:
                    run.font.underline = True
                    rgb = self._map_to_rgb(color_name, RGBColor)
                    if rgb is not None:
                        run.font.color.rgb = rgb

        # Save to memory and return base64
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return b64encode(buf.getvalue()).decode("ascii")

    def visualize(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
        output_format: str = "docx",
    ) -> str:
        """
        Convenience: build + render.

        Returns:
        - base64-encoded DOCX (str)
        """
        spec = self.build(cas, start=start, end=end)
        return self.render(spec, output_format=output_format)

    # ------------- helpers -------------

    @staticmethod
    def _check_overlap(sorted_spans: list[dict[str, Any]]) -> bool:
        """
        Detect overlap in a list of spans sorted by (start, end).
        Returns True if any span starts before the previous one ended.
        """
        prev_end = -1
        for s in sorted_spans:
            if s["start"] < prev_end:
                return True
            prev_end = max(prev_end, s["end"])
        return False

    @staticmethod
    def _map_to_highlight(color_name: str | None, WD_COLOR_INDEX) -> Any:
        """
        Map a CSS-like color name to Word's WD_COLOR_INDEX (limited highlight palette).

        Fallback:
        - if color_name is None or unknown, returns WD_COLOR_INDEX.YELLOW
        """
        if color_name is None:
            return WD_COLOR_INDEX.YELLOW
        name = color_name.lower()
        mapping = {
            "yellow": WD_COLOR_INDEX.YELLOW,
            "lightyellow": WD_COLOR_INDEX.YELLOW,
            "orange": WD_COLOR_INDEX.DARK_YELLOW,
            "orangered": WD_COLOR_INDEX.RED,
            "red": WD_COLOR_INDEX.RED,
            "darkred": WD_COLOR_INDEX.DARK_RED,
            "lightgreen": WD_COLOR_INDEX.BRIGHT_GREEN,
            "palegreen": WD_COLOR_INDEX.BRIGHT_GREEN,
            "green": WD_COLOR_INDEX.GREEN,
            "mediumseagreen": WD_COLOR_INDEX.GREEN,
            "turquoise": WD_COLOR_INDEX.TURQUOISE,
            "paleturquoise": WD_COLOR_INDEX.TURQUOISE,
            "skyblue": WD_COLOR_INDEX.TURQUOISE,
            "steelblue": WD_COLOR_INDEX.BLUE,
            "blue": WD_COLOR_INDEX.BLUE,
            "plum": WD_COLOR_INDEX.PINK,
            "pink": WD_COLOR_INDEX.PINK,
            "violet": WD_COLOR_INDEX.VIOLET,
            "mediumpurple": WD_COLOR_INDEX.VIOLET,
            "silver": WD_COLOR_INDEX.GRAY_25,
            "gray": WD_COLOR_INDEX.GRAY_50,
            "black": WD_COLOR_INDEX.BLACK,
            "navajowhite": WD_COLOR_INDEX.DARK_YELLOW,
            "rosybrown": WD_COLOR_INDEX.DARK_RED,
        }
        return mapping.get(name, WD_COLOR_INDEX.YELLOW)

    @staticmethod
    def _map_to_rgb(color_name: str | None, RGBColor) -> Optional[Any]:
        """
        Map a CSS-like color name to an RGBColor for text foreground (underline mode).

        Returns:
        - RGBColor or None if unknown or color_name is None.
        """
        if color_name is None:
            return None
        name = color_name.lower()
        hexmap = {
            "lightgreen": "90EE90",
            "orangered": "FF4500",
            "orange": "FFA500",
            "plum": "DDA0DD",
            "palegreen": "98FB98",
            "mediumseagreen": "3CB371",
            "steelblue": "4682B4",
            "skyblue": "87CEEB",
            "navajowhite": "FFDEAD",
            "mediumpurple": "9370DB",
            "rosybrown": "BC8F8F",
            "silver": "C0C0C0",
            "gray": "808080",
            "paleturquoise": "AFEEEE",
            "blue": "0000FF",
            "red": "FF0000",
            "green": "008000",
            "yellow": "FFFF00",
            "violet": "EE82EE",
            "pink": "FFC0CB",
            "turquoise": "40E0D0",
            "black": "000000",
        }
        hx = hexmap.get(name)
        if hx is None:
            return None
        r = int(hx[0:2], 16)
        g = int(hx[2:4], 16)
        b = int(hx[4:6], 16)
        return RGBColor(r, g, b)

# ---------- heatmap visualizer ----------

class HeatmapVisualizer(Visualizer):
    """
    Token-level heat map showing where configured annotations occur, one cell per token.

    Visual semantics:
    - 'binary': a token-cell is filled if any configured annotation overlaps the token (annotated color).
    - 'density': a token-cell opacity scales with the number of overlapping annotations (annotated color).
    - Non-annotated tokens get a contrasting fill color to make sentence length and structure visible.

    Row modes:
    - 'flat': grid wraps after max_cols (single token stream).
    - 'sentence': one row per sentence; each row ends at its sentence boundary (rigged layout).

    Grid geometry:
    - flat: cols = min(token_count, max_cols); rows = ceil(token_count / cols).
    - sentence: cols = max token count across all sentences in range; rows = number of sentences.

    Parameters:
    - ts: TypeSystem | Path | str
    - token_type: fully qualified CAS type name of tokens
      e.g., "de.tudarmstadt.ukp.dkpro.core.api.segmentation.type.Token"
    - sentence_type: sentence annotation type; defaults to short "Sentence" with DKPro fallback.
    - types: list[str] | None
      CAS annotation types to include in the heat map (use add_type(...) as alternative).
    - mode: 'binary' | 'density' (default: 'binary')
    - row_mode: 'sentence' | 'flat' (default: 'sentence')
    - color_hex: '#RRGGBB' annotated cell base color (default: '#FF4500' orangered)
    - unannotated_hex: '#RRGGBB' non-annotated cell color (default: complement of color_hex)
    - background_hex: '#RRGGBB' background (default: '#FFFFFF')
    - ann_alpha_max: max alpha for annotated cells (default: 0.9)
    - unann_alpha: alpha for non-annotated cells (default: 0.25)
    - max_cols: maximum logical columns for 'flat' mode (default: 40)
    - cell_px: CSS pixels per logical cell (default: 2)
    - width_px: optional CSS target width (default: None). If provided, pixels-per-cell ≈ floor(width_px / cols).
    - page: wrap fragment into a full HTML page (UTF‑8) if True (default: False)
    - strict: raise when no tokens or no coverage are found (default: True)

    Range semantics:
    - Only tokens fully inside [start, end) are part of the grid.
    - An annotation overlaps a token if their half-open intervals intersect:
      token [t_b, t_e) overlaps fs [f_b, f_e) iff t_b < f_e and f_b < t_e.
    """

    def __init__(
        self,
        ts: TypeSystem | Path | str,
        *,
        token_type: str,
        sentence_type: str = "Sentence",
        types: list[str] | None = None,
        mode: str = "binary",
        row_mode: str = "sentence",
        color_hex: str = "#FF4500",
        unannotated_hex: str | None = None,
        background_hex: str = "#FFFFFF",
        ann_alpha_max: float = 0.9,
        unann_alpha: float = 0.25,
        max_cols: int = 40,
        cell_px: int = 2,
        width_px: int | None = None,
        page: bool = False,
        strict: bool = True,
    ):
        # Accept Path by converting to str for base init
        super().__init__(ts if isinstance(ts, TypeSystem) else str(ts))

        if not token_type:
            raise VisualizerException("HeatmapVisualizer: token_type must be provided")
        self._token_type = token_type

        if mode not in ("binary", "density"):
            raise VisualizerException("HeatmapVisualizer: mode must be 'binary' or 'density'")
        self._mode = mode

        if row_mode not in ("sentence", "flat"):
            raise VisualizerException("HeatmapVisualizer: row_mode must be 'sentence' or 'flat'")
        self._row_mode = row_mode

        if max_cols <= 0:
            raise VisualizerException("HeatmapVisualizer: max_cols must be > 0")
        self._max_cols = max_cols

        if cell_px <= 0:
            raise VisualizerException("HeatmapVisualizer: cell_px must be > 0")
        self._cell_px = cell_px

        if width_px is not None and width_px <= 0:
            raise VisualizerException("HeatmapVisualizer: width_px must be > 0 if provided")
        self._width_px = width_px

        self._sentence_type = sentence_type
        self._sentence_fallback = "de.tudarmstadt.ukp.dkpro.core.api.segmentation.type.Sentence"

        self._color_hex = color_hex
        self._background_hex = background_hex
        # auto-compute a contrasting color if not provided (simple complement)
        self._unannotated_hex = unannotated_hex or self._hex_complement(color_hex)

        if not (0.0 < ann_alpha_max <= 1.0):
            raise VisualizerException("HeatmapVisualizer: ann_alpha_max must be in (0, 1]")
        if not (0.0 < unann_alpha <= 1.0):
            raise VisualizerException("HeatmapVisualizer: unann_alpha must be in (0, 1]")
        self._ann_alpha_max = ann_alpha_max
        self._unann_alpha = unann_alpha

        self._page = page
        self._strict = strict

        if types:
            for t in types:
                self.add_type(t)

    # ---------- build / render / visualize ----------

    def build(self, cas: Cas, *, start: int = 0, end: int = -1) -> Dict[str, Any]:
        text_len = len(cas.sofa_string)
        if end == -1:
            end = text_len
        if start < 0 or end < 0 or start > end or end > text_len:
            raise VisualizerException(
                f"HeatmapVisualizer: invalid range [{start}, {end}] for document length {text_len}"
            )

        type_names = self.list_types()
        if not type_names:
            raise VisualizerException("HeatmapVisualizer: no annotation types configured. Use add_type(...) first.")

        if self._row_mode == "flat":
            tokens = [tok for tok in cas.select(self._token_type) if tok.begin >= start and tok.end <= end]
            if not tokens:
                raise VisualizerException(
                    f"HeatmapVisualizer: no tokens of type {self._token_type} found inside range [{start}, {end}]."
                )

            n = len(tokens)
            begins = [t.begin for t in tokens]
            ends = [t.end for t in tokens]

            diff = [0] * (n + 1)
            have_any_overlap = False

            for type_name in type_names:
                for fs in cas.select(type_name):
                    fb = max(fs.begin, start)
                    fe = min(fs.end, end)
                    if fb >= fe:
                        continue
                    li = bisect_right(ends, fb)
                    ri_excl = bisect_left(begins, fe)
                    if li < ri_excl:
                        diff[li] += 1
                        diff[ri_excl] -= 1
                        have_any_overlap = True

            coverage = [0] * n
            cur = 0
            max_count = 0
            for i in range(n):
                cur += diff[i]
                coverage[i] = cur
                if cur > max_count:
                    max_count = cur

            if self._strict and not have_any_overlap:
                raise VisualizerException(
                    f"HeatmapVisualizer: no overlaps found for types {type_names} over tokens in range [{start}, {end}]."
                )
            if self._strict and max_count == 0:
                raise VisualizerException(
                    f"HeatmapVisualizer: nothing to render (no covered tokens) in range [{start}, {end}]."
                )

            cols = min(n, self._max_cols)
            rows = (n + cols - 1) // cols  # ceil
            css_cell_px, css_width_px = self._compute_css_cell_and_width(cols)
            css_height_px = rows * css_cell_px

            r, g, b = self._hex_to_rgb(self._color_hex)
            ur, ug, ub = self._hex_to_rgb(self._unannotated_hex)
            br, bg, bb = self._hex_to_rgb(self._background_hex)

            return {
                "fill_mode": self._mode,
                "row_mode": "flat",
                "color": [r, g, b],
                "unannotated_color": [ur, ug, ub],
                "background": [br, bg, bb],
                "ann_alpha_max": self._ann_alpha_max,
                "unann_alpha": self._unann_alpha,
                "cols": cols,
                "rows": rows,
                "coverage": coverage,   # 1-D coverage
                "max_count": max_count,
                "css_cell_px": css_cell_px,
                "css_width_px": css_width_px,
                "css_height_px": css_height_px,
            }

        # row_mode == 'sentence'
        sentences = list(cas.select(self._sentence_type))
        if not sentences and "." not in self._sentence_type:
            sentences = list(cas.select(self._sentence_fallback))
        sentences = [s for s in sentences if s.begin >= start and s.end <= end]
        if not sentences:
            raise VisualizerException(
                f"HeatmapVisualizer: no sentences of type {self._sentence_type} found inside range [{start}, {end}]."
            )

        sent_tokens: list[list[FeatureStructure]] = []
        tokens_per_row: list[int] = []
        for s in sentences:
            toks = list(cas.select_covered(self._token_type, covering_annotation=s))
            toks = [t for t in toks if t.begin >= start and t.end <= end]
            sent_tokens.append(toks)
            tokens_per_row.append(len(toks))

        if all(n == 0 for n in tokens_per_row):
            raise VisualizerException(
                f"HeatmapVisualizer: sentences contain no tokens of type {self._token_type} inside range [{start}, {end}]."
            )

        have_any_overlap = False
        max_count = 0
        max_cols_for_sentences = max(tokens_per_row)
        rows = len(sent_tokens)
        cols = max_cols_for_sentences

        coverage_rows: list[list[int]] = []
        for s_idx, s in enumerate(sentences):
            toks = sent_tokens[s_idx]
            n = len(toks)
            begins = [t.begin for t in toks]
            ends = [t.end for t in toks]
            diff = [0] * (n + 1)

            for type_name in type_names:
                for fs in cas.select_covered(type_name, covering_annotation=s):
                    fb = max(fs.begin, start)
                    fe = min(fs.end, end)
                    if fb >= fe:
                        continue
                    li = bisect_right(ends, fb)
                    ri_excl = bisect_left(begins, fe)
                    if li < ri_excl:
                        diff[li] += 1
                        diff[ri_excl] -= 1
                        have_any_overlap = True

            row_cov = [0] * n
            cur = 0
            for i in range(n):
                cur += diff[i]
                row_cov[i] = cur
                if cur > max_count:
                    max_count = cur

            # Pad with zeros to max_cols_for_sentences (padded cells represent "no token")
            if n < max_cols_for_sentences:
                row_cov = row_cov + [0] * (max_cols_for_sentences - n)

            coverage_rows.append(row_cov)

        if self._strict and not have_any_overlap:
            raise VisualizerException(
                f"HeatmapVisualizer: no overlaps found for types {type_names} over sentences in range [{start}, {end}]."
            )
        if self._strict and max_count == 0:
            raise VisualizerException(
                f"HeatmapVisualizer: nothing to render (no covered tokens) in range [{start}, {end}]."
            )

        coverage = [c for row in coverage_rows for c in row]

        css_cell_px, css_width_px = self._compute_css_cell_and_width(cols)
        css_height_px = rows * css_cell_px

        r, g, b = self._hex_to_rgb(self._color_hex)
        ur, ug, ub = self._hex_to_rgb(self._unannotated_hex)
        br, bg, bb = self._hex_to_rgb(self._background_hex)

        return {
            "fill_mode": self._mode,
            "row_mode": "sentence",
            "color": [r, g, b],
            "unannotated_color": [ur, ug, ub],
            "background": [br, bg, bb],
            "ann_alpha_max": self._ann_alpha_max,
            "unann_alpha": self._unann_alpha,
            "cols": cols,
            "rows": rows,
            "coverage": coverage,               # flattened row-major
            "tokens_per_row": tokens_per_row,   # differentiate padded cells vs real tokens
            "max_count": max_count,
            "css_cell_px": css_cell_px,
            "css_width_px": css_width_px,
            "css_height_px": css_height_px,
        }

    def render(self, spec: Dict[str, Any], *, output_format: str = "html") -> str:
        fmt = output_format.lower()
        if fmt != "html":
            raise VisualizerException("HeatmapVisualizer supports only 'html' output_format")
        frag = self._render_canvas_fragment(spec)
        return self._wrap_html_page(frag, "HeatmapVisualizer") if self._page else frag

    def visualize(self, cas: Cas, *, start: int = 0, end: int = -1, output_format: str = "html") -> str:
        spec = self.build(cas, start=start, end=end)
        return self.render(spec, output_format=output_format)

    # ---------- internal helpers ----------

    def _compute_css_cell_and_width(self, cols: int) -> tuple[int, int]:
        """
        Compute CSS pixels per logical cell and total width.
        - If width_px is None: use cell_px exactly.
        - If width_px is provided: respect requested cell_px where possible,
        but reduce it if necessary so that cols * css_cell_px <= width_px.
        """
        if self._width_px is None:
            css_cell_px = self._cell_px
            css_width_px = cols * css_cell_px
        else:
            # Width-imposed cell size (integer pixels per cell)
            imposed = self._width_px // max(cols, 1)
            if imposed < 1:
                imposed = 1
            # Use the smaller of requested and imposed to respect the width constraint
            css_cell_px = min(self._cell_px, imposed)
            css_width_px = cols * css_cell_px
        return css_cell_px, css_width_px

    @staticmethod
    def _hex_to_rgb(hx: str) -> tuple[int, int, int]:
        s = hx.strip().lstrip("#")
        if len(s) != 6:
            return (255, 69, 0)  # default orangered
        return (int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))

    @staticmethod
    def _hex_complement(hx: str) -> str:
        s = hx.strip().lstrip("#")
        if len(s) != 6:
            return "#00B2FF"  # some contrasting default (azure)
        r = 255 - int(s[0:2], 16)
        g = 255 - int(s[2:4], 16)
        b = 255 - int(s[4:6], 16)
        return f"#{r:02X}{g:02X}{b:02X}"

    def _render_canvas_fragment(self, spec: Dict[str, Any]) -> str:
        import json

        fill_mode = spec["fill_mode"]
        color = spec["color"]
        unann = spec["unannotated_color"]
        bg = spec["background"]
        ann_alpha_max = float(spec["ann_alpha_max"])
        unann_alpha = float(spec["unann_alpha"])
        cols = int(spec["cols"])
        rows = int(spec["rows"])
        coverage: list[int] = spec["coverage"]
        tokens_per_row: list[int] | None = spec.get("tokens_per_row")
        max_count = int(spec["max_count"])
        css_width_px = int(spec["css_width_px"])
        css_height_px = int(spec["css_height_px"])

        data = {
            "fillMode": fill_mode,
            "annColor": color,
            "unannColor": unann,
            "background": bg,
            "annAlphaMax": ann_alpha_max,
            "unannAlpha": unann_alpha,
            "cols": cols,
            "rows": rows,
            "coverage": coverage,
            "tokensPerRow": tokens_per_row,
            "maxCount": max_count,
        }
        js_data = json.dumps(data)

        return f"""
<div style="display:inline-block;border:1px solid #ccc;padding:6px;background:#f9f9f9;">
  <div style="font:14px/1.4 system-ui, -apple-system, Segoe UI, Roboto, Ubuntu, Cantarell, Arial;">
    <strong>Token heat map</strong><br>cols={cols}, rows={rows}, mode={fill_mode}
  </div>
  <canvas id="hm_canvas" width="{cols}" height="{rows}"
          style="width:{css_width_px}px;height:{css_height_px}px;image-rendering:pixelated;border:1px solid #ddd;"></canvas>
</div>
<script>
(function() {{
  const data = {js_data};
  const cvs = document.getElementById('hm_canvas');
  const ctx = cvs.getContext('2d');

  const cols = data.cols;
  const rows = data.rows;
  const cov = data.coverage;
  const maxC = data.maxCount;
  const ann = data.annColor;    // [r,g,b]
  const unann = data.unannColor; // [r,g,b]
  const bg = data.background;
  const annA = data.annAlphaMax;
  const unannA = data.unannAlpha;
  const tpr = data.tokensPerRow; // null in flat mode

  // Fill background
  ctx.fillStyle = 'rgb(' + bg[0] + ',' + bg[1] + ',' + bg[2] + ')';
  ctx.fillRect(0, 0, cvs.width, cvs.height);

  // Draw cells: one logical pixel per token; CSS scales via style width/height
  let k = 0;
  for (let y = 0; y < rows; y++) {{
    const rowLen = tpr ? tpr[y] : cols; // real tokens in this row
    for (let x = 0; x < cols; x++) {{
      if (k >= cov.length) break;
      const c = cov[k];
      const isRealToken = x < rowLen;
      if (isRealToken) {{
        if (data.fillMode === 'binary') {{
          if (c > 0) {{
            ctx.fillStyle = 'rgba(' + ann[0] + ',' + ann[1] + ',' + ann[2] + ',' + annA + ')';
            ctx.fillRect(x, y, 1, 1);
          }} else {{
            ctx.fillStyle = 'rgba(' + unann[0] + ',' + unann[1] + ',' + unann[2] + ',' + unannA + ')';
            ctx.fillRect(x, y, 1, 1);
          }}
        }} else {{
          if (c > 0 && maxC > 0) {{
            const alpha = Math.min(annA, 0.1 + (annA - 0.1) * (c / maxC));
            ctx.fillStyle = 'rgba(' + ann[0] + ',' + ann[1] + ',' + ann[2] + ',' + alpha + ')';
            ctx.fillRect(x, y, 1, 1);
          }} else {{
            ctx.fillStyle = 'rgba(' + unann[0] + ',' + unann[1] + ',' + unann[2] + ',' + unannA + ')';
            ctx.fillRect(x, y, 1, 1);
          }}
        }}
      }}
      k++;
    }}
  }}
}})();
</script>
""".strip()