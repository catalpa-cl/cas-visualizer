from __future__ import annotations

import re
from bisect import bisect_left, bisect_right
from pathlib import Path
from typing import Any, List

import pandas as pd
from cassis import Cas
from cassis.typesystem import FeatureStructure
from docx.oxml.ns import qn
from docx.shared import RGBColor
from spacy.displacy import EntityRenderer, SpanRenderer

from cas_visualizer._base import Visualizer, VisualizerException

# ---------- Span visualizer ----------

class SpacySpanVisualizer(Visualizer):
    """
    Span visualization using spaCy’s displaCy.

    Modes:
    - HIGHLIGHT: render entity-like colored spans via EntityRenderer
    - UNDERLINE: render underlined spans via SpanRenderer (token-based)

    Notes:
    - displaCy’s color map is keyed by label strings only, so different types with the same label share colors.
      To avoid collisions, prefer unique labels across types.
    """
    HIGHLIGHT: str = "HIGHLIGHT"
    UNDERLINE: str = "UNDERLINE"

    def __init__(
        self,
        ts: str | Path | TypeSystem,
        span_type: str | None = None,
        types: list[str] | None = None,
        *,
        page: bool = False,
        strict: bool = True,
    ):
        super().__init__(ts)
        self._span_types: list[str] = [SpacySpanVisualizer.HIGHLIGHT, SpacySpanVisualizer.UNDERLINE]
        self._selected_span_type = SpacySpanVisualizer.UNDERLINE
        if span_type is not None:
            self.selected_span_type = span_type
        self._allow_highlight_overlap = False
        self._page = page
        self._strict = strict
        if types is not None:
            for type_name in types:
                self.add_type(type_name)

    @property
    def selected_span_type(self) -> str:
        """Current visualization mode: 'HIGHLIGHT' or 'UNDERLINE'."""
        return self._selected_span_type

    @selected_span_type.setter
    def selected_span_type(self, value: str) -> None:
        if value not in self._span_types:
            raise VisualizerException(f"Invalid span type: {value}. Expected one of {self._span_types}")
        self._selected_span_type = value

    @property
    def allow_highlight_overlap(self) -> bool:
        """Whether overlapping highlights are allowed (HIGHLIGHT mode)."""
        return self._allow_highlight_overlap

    @allow_highlight_overlap.setter
    def allow_highlight_overlap(self, value: bool):
        self._allow_highlight_overlap = value

    # ------------- build / render / visualize -------------

    def build(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
    ) -> dict[str, Any]:
        """
        Build a displaCy spec for the selected span visualization.

        Returns:
        - HIGHLIGHT: {'mode': 'HIGHLIGHT', 'text': str, 'ents': list[dict], 'colors': dict[label, color]}
        - UNDERLINE: {'mode': 'UNDERLINE', 'token_texts': list[str], 'spans': list[dict], 'colors': dict[label, color]}

        Errors:
        - VisualizerException if overlapping highlights are detected (and overlap not allowed) or no content in strict mode.
        """
        if end == -1:
            end = len(cas.sofa_string)

        if self.selected_span_type == SpacySpanVisualizer.HIGHLIGHT:
            tmp_ents: list[dict[str, Any]] = []
            labels_to_colors: dict[str, str] = {}
            for annotation_type in self.list_types():
                for fs in cas.select(annotation_type):
                    if fs.begin < start or fs.end > end:
                        continue
                    label = self.resolve_label(fs, annotation_type)
                    color = self.resolve_color(annotation_type, label)
                    if color:
                        tmp_ents.append({"start": fs.begin, "end": fs.end, "label": label})
                        labels_to_colors[label] = color

            sorted_ents = sorted(tmp_ents, key=lambda x: (x["start"], x["end"]))
            if not self._allow_highlight_overlap and self._check_overlap(sorted_ents):
                raise VisualizerException(
                    "Highlighted annotations are overlapping. Choose a different set of annotations "
                    "or set allow_highlight_overlap=True."
                )
            if self._strict and not sorted_ents:
                raise VisualizerException(
                    f"No entities found for configured types {self.list_types()} in range [{start}, {end}]."
                )

            return {
                "mode": SpacySpanVisualizer.HIGHLIGHT,
                "text": cas.sofa_string,
                "ents": sorted_ents,
                "colors": labels_to_colors,
            }

        elif self.selected_span_type == SpacySpanVisualizer.UNDERLINE:
            selected_annotations = [
                item for typeclass in self.list_types() for item in cas.select(typeclass)
                if item.begin >= start and item.end <= end
            ]
            if self._strict and not selected_annotations:
                raise VisualizerException(
                    f"No spans found for configured types {self.list_types()} in range [{start}, {end}]."
                )

            tmp_tokens = self._split_text_for_spans(cas.sofa_string, selected_annotations)
            tmp_token_texts = [_["text"] for _ in sorted(tmp_tokens, key=lambda t: t["start"])]

            tmp_spans: list[dict[str, Any]] = []
            labels_to_colors: dict[str, str] = {}
            for annotation_type in self.list_types():
                for tmp_span in self._create_spans(
                    cas=cas, cas_sofa_tokens=tmp_tokens, annotation_type=annotation_type, start=start, end=end
                ):
                    label = tmp_span["label"]
                    color = self.resolve_color(annotation_type, label)
                    if color is not None:
                        labels_to_colors[label] = color
                        tmp_spans.append(tmp_span)

            tmp_spans.sort(key=lambda x: x["start"])
            if self._strict and not tmp_spans:
                raise VisualizerException(
                    f"No underline spans could be created for configured types {self.list_types()} in range [{start}, {end}]."
                )

            return {
                "mode": SpacySpanVisualizer.UNDERLINE,
                "token_texts": tmp_token_texts,
                "spans": tmp_spans,
                "colors": labels_to_colors,
            }

        else:
            raise VisualizerException("Invalid span type")

    def render(
        self,
        spec: dict[str, Any],
        *,
        output_format: str = "html",
    ) -> str:
        """
        Render the built spec using spaCy’s displaCy.

        Supported formats:
        - 'html' only (returns a string). Set page=True to get a full HTML document.

        Errors:
        - VisualizerException on unsupported format or invalid spec mode.
        """
        fmt = output_format.lower()
        if fmt != "html":
            raise VisualizerException("SpacySpanVisualizer supports only 'html' output_format")

        mode = spec.get("mode")
        if mode == SpacySpanVisualizer.HIGHLIGHT:
            frag = EntityRenderer({"colors": spec["colors"]}).render_ents(spec["text"], spec["ents"], "")
        elif mode == SpacySpanVisualizer.UNDERLINE:
            frag = SpanRenderer({"colors": spec["colors"]}).render_spans(spec["token_texts"], spec["spans"], "")
        else:
            raise VisualizerException("Invalid spec: missing or unknown mode")

        return self._wrap_html_page(frag, "SpanVisualizer") if self._page else frag

    def visualize(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
        output_format: str = "html",
    ) -> str:
        """Convenience: build + render (returns str)."""
        spec = self.build(cas, start=start, end=end)
        return self.render(spec, output_format=output_format)

    # ------------- private helpers -------------

    @staticmethod
    def _check_overlap(sorted_ents: list[dict[str, Any]]) -> bool:
        """
        Detect overlap in a list of spans sorted by (start, end).
        Returns True if any span starts before the previous one ended.
        """
        prev_end = -1
        for ent in sorted_ents:
            if ent["start"] < prev_end:
                return True
            prev_end = max(prev_end, ent["end"])
        return False

    @staticmethod
    def _split_text_for_spans(
        cas_sofa_string: str,
        feature_structures: list[FeatureStructure],
    ) -> list[dict[str, int | str]]:
        """
        Split text at annotation boundaries and after whitespace to build token-like chunks for SpanRenderer.
        Returns list of dicts: {start, end, text}.
        """
        cutting_points = {fs.begin for fs in feature_structures} | {fs.end for fs in feature_structures}
        cutting_points |= {i + 1 for i, ch in enumerate(cas_sofa_string) if ch.isspace()}

        points = sorted(cutting_points | {0, len(cas_sofa_string)})

        tokens: list[dict[str, int | str]] = []
        for i in range(len(points) - 1):
            start, end = points[i], points[i + 1]
            if start == end:
                continue
            tokens.append({"start": start, "end": end, "text": cas_sofa_string[start:end]})

        return tokens

    def _create_spans(
        self,
        cas: Cas,
        cas_sofa_tokens: list[dict[str, int | str]],
        annotation_type: str,
        *,
        start: int,
        end: int,
    ) -> list[dict[str, int | str | None]]:
        """
        Project annotation character offsets onto token indices for the SpanRenderer.

        Returns dicts with:
        - start/end (char offsets),
        - start_token/end_token (token indices),
        - label (resolved via type configuration).
        """
        start_idx = {tok["start"]: i for i, tok in enumerate(cas_sofa_tokens)}
        end_idx_next = {tok["end"]: i + 1 for i, tok in enumerate(cas_sofa_tokens)}  # exclusive end

        spans: list[dict[str, int | str | None]] = []
        for fs in cas.select(annotation_type):
            if fs.begin < start or fs.end > end:
                continue
            try:
                start_tok = start_idx[fs.begin]
                end_tok = end_idx_next[fs.end]
            except KeyError as e:
                raise VisualizerException(
                    f"Annotation [{fs.begin}, {fs.end}] of type {annotation_type} "
                    f"does not align with token boundaries."
                ) from e

            spans.append(
                {
                    "start": fs.begin,
                    "end": fs.end,
                    "start_token": start_tok,
                    "end_token": end_tok,
                    "label": self.resolve_label(fs, annotation_type),
                }
            )
        return spans

# ---------- DOCX span visualizer (Word document with highlighting/underline) ----------

class DocxSpanVisualizer(Visualizer):
    """
    Span visualization that outputs a Word (.docx) document with text and highlighting/underline.

    Modes:
    - HIGHLIGHT: highlight spans with background colors (limited to Word's WD_COLOR_INDEX palette).
    - UNDERLINE: underline spans and optionally set the text color.

    Contract:
    - build: similar to SpacySpanVisualizer; gathers spans within [start, end].
      Returns a spec dict: {'mode': 'HIGHLIGHT'|'UNDERLINE', 'text': str, 'spans': list[dict]}
        Each span dict: {'start': int, 'end': int, 'label': str|None, 'color': str|None}
    - render: supports 'docx' only; returns base64-encoded DOCX bytes as a string.
    - visualize: build + render (returns the base64-encoded DOCX string).

    Strictness and overlap:
    - If strict=True and no spans are found, a VisualizerException is raised.
    - In HIGHLIGHT mode with allow_highlight_overlap=False, overlapping spans raise VisualizerException.
      If allow_highlight_overlap=True, overlapping regions will be formatted with the most recently begun span.

    Dependencies:
    - Requires `python-docx` (pip install python-docx).
    """

    HIGHLIGHT: str = "HIGHLIGHT"
    UNDERLINE: str = "UNDERLINE"

    def __init__(
        self,
        ts: str | Path | TypeSystem,
        *,
        mode: str = HIGHLIGHT,
        types: list[str] | None = None,
        allow_highlight_overlap: bool = False,
        strict: bool = True,
    ):
        super().__init__(ts)
        if mode not in (DocxSpanVisualizer.HIGHLIGHT, DocxSpanVisualizer.UNDERLINE):
            raise VisualizerException(
                f"Invalid mode: {mode}. Expected one of [{DocxSpanVisualizer.HIGHLIGHT}, {DocxSpanVisualizer.UNDERLINE}]"
            )
        self._mode = mode
        self._allow_highlight_overlap = allow_highlight_overlap
        self._strict = strict
        if types is not None:
            for type_name in types:
                self.add_type(type_name)

    # ------------- build / render / visualize -------------

    def build(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
    ) -> dict[str, Any]:
        """
        Build a spec for DOCX rendering.

        Returns:
        - {'mode': 'HIGHLIGHT'|'UNDERLINE', 'text': str, 'spans': list[dict]}
          where each span dict has: {'start': int, 'end': int, 'label': str|None, 'color': str|None}

        Errors:
        - VisualizerException if overlapping highlights are detected (and overlap not allowed) or no content in strict mode.
        """
        if end == -1:
            end = len(cas.sofa_string)

        spans: list[dict[str, Any]] = []
        for annotation_type in self.list_types():
            for fs in cas.select(annotation_type):
                if fs.begin < start or fs.end > end:
                    continue
                label = self.resolve_label(fs, annotation_type)
                color = self.resolve_color(annotation_type, label)
                spans.append({"start": fs.begin, "end": fs.end, "label": label, "color": color})

        spans.sort(key=lambda s: (s["start"], s["end"]))

        if self._mode == DocxSpanVisualizer.HIGHLIGHT and not self._allow_highlight_overlap:
            if self._check_overlap(spans):
                raise VisualizerException(
                    "Highlighted annotations are overlapping. Choose a different set of annotations "
                    "or enable allow_highlight_overlap=True."
                )

        if self._strict and not spans:
            raise VisualizerException(
                f"No spans found for configured types {self.list_types()} in range [{start}, {end}]."
            )

        return {"mode": self._mode, "text": cas.sofa_string, "spans": spans}

    def render(
        self,
        spec: dict[str, Any],
        *,
        output_format: str = "docx",
    ) -> str:
        """
        Render the built spec to an in-memory .docx and return base64-encoded bytes.

        Supported formats:
        - 'docx' only

        Returns:
        - base64-encoded DOCX (str)

        Errors:
        - VisualizerException on unsupported format, invalid spec mode, or missing python-docx.
        """
        fmt = output_format.lower()
        if fmt != "docx":
            raise VisualizerException("DocxSpanVisualizer supports only 'docx' output_format")

        try:
            from io import BytesIO
            from base64 import b64encode
            from docx import Document  # python-docx
            from docx.enum.text import WD_COLOR_INDEX
            from docx.shared import RGBColor
        except ImportError as e:
            raise VisualizerException("python-docx is required for DOCX rendering (pip install python-docx).") from e

        mode = spec.get("mode")
        if mode not in (DocxSpanVisualizer.HIGHLIGHT, DocxSpanVisualizer.UNDERLINE):
            raise VisualizerException("Invalid spec: missing or unknown mode")

        text: str = spec["text"]
        spans: list[dict[str, Any]] = spec["spans"]

        # Prepare events for a sweep-line over the text to create formatted runs
        events: list[tuple[int, str, dict[str, Any]]] = []  # (pos, 'begin'|'end', span)
        for s in spans:
            events.append((s["start"], "begin", s))
            events.append((s["end"], "end", s))
        # Process end events before begin events at the same position to avoid zero-width segments
        events.sort(key=lambda t: (t[0], 0 if t[1] == "end" else 1))

        doc = Document()
        para = doc.add_paragraph()

        active: list[dict[str, Any]] = []
        prev_pos = 0

        for pos, kind, span in events:
            if pos > prev_pos:
                segment_text = text[prev_pos:pos]
                run = para.add_run(segment_text)
                # Apply style based on current active spans
                style_span = active[-1] if active else None  # last-opened span wins for overlap
                if style_span is not None:
                    color_name = style_span.get("color")
                    if mode == DocxSpanVisualizer.HIGHLIGHT:
                        run.font.highlight_color = self._map_to_highlight(color_name, WD_COLOR_INDEX)  # type: ignore[attr-defined]
                    elif mode == DocxSpanVisualizer.UNDERLINE:
                        run.font.underline = True
                        rgb = self._map_to_rgb(color_name, RGBColor)
                        if rgb is not None:
                            run.font.color.rgb = rgb
                prev_pos = pos

            # Update active set
            if kind == "begin":
                active.append(span)
            else:  # 'end'
                # remove the first matching span by identity or by (start,end,label)
                for i, a in enumerate(active):
                    if a is span or (a["start"], a["end"], a.get("label")) == (span["start"], span["end"], span.get("label")):
                        active.pop(i)
                        break

        # Tail segment
        if prev_pos < len(text):
            run = para.add_run(text[prev_pos:])
            style_span = active[-1] if active else None
            if style_span is not None:
                color_name = style_span.get("color")
                if mode == DocxSpanVisualizer.HIGHLIGHT:
                    run.font.highlight_color = self._map_to_highlight(color_name, WD_COLOR_INDEX)  # type: ignore[attr-defined]
                elif mode == DocxSpanVisualizer.UNDERLINE:
                    run.font.underline = True
                    rgb = self._map_to_rgb(color_name, RGBColor)
                    if rgb is not None:
                        run.font.color.rgb = rgb

        # Save to memory and return base64
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return b64encode(buf.getvalue()).decode("ascii")

    def visualize(
        self,
        cas: Cas,
        *,
        start: int = 0,
        end: int = -1,
        output_format: str = "docx",
    ) -> str:
        """
        Convenience: build + render.

        Returns:
        - base64-encoded DOCX (str)
        """
        spec = self.build(cas, start=start, end=end)
        return self.render(spec, output_format=output_format)

    # ------------- helpers -------------

    @staticmethod
    def _check_overlap(sorted_spans: list[dict[str, Any]]) -> bool:
        """
        Detect overlap in a list of spans sorted by (start, end).
        Returns True if any span starts before the previous one ended.
        """
        prev_end = -1
        for s in sorted_spans:
            if s["start"] < prev_end:
                return True
            prev_end = max(prev_end, s["end"])
        return False

    @staticmethod
    def _map_to_highlight(color_name: str | None, WD_COLOR_INDEX) -> Any:
        """
        Map a CSS-like color name to Word's WD_COLOR_INDEX (limited highlight palette).

        Fallback:
        - if color_name is None or unknown, returns WD_COLOR_INDEX.YELLOW
        """
        if color_name is None:
            return WD_COLOR_INDEX.YELLOW
        name = color_name.lower()
        mapping = {
            "yellow": WD_COLOR_INDEX.YELLOW,
            "lightyellow": WD_COLOR_INDEX.YELLOW,
            "orange": WD_COLOR_INDEX.DARK_YELLOW,
            "orangered": WD_COLOR_INDEX.RED,
            "red": WD_COLOR_INDEX.RED,
            "darkred": WD_COLOR_INDEX.DARK_RED,
            "lightgreen": WD_COLOR_INDEX.BRIGHT_GREEN,
            "palegreen": WD_COLOR_INDEX.BRIGHT_GREEN,
            "green": WD_COLOR_INDEX.GREEN,
            "mediumseagreen": WD_COLOR_INDEX.GREEN,
            "turquoise": WD_COLOR_INDEX.TURQUOISE,
            "paleturquoise": WD_COLOR_INDEX.TURQUOISE,
            "skyblue": WD_COLOR_INDEX.TURQUOISE,
            "steelblue": WD_COLOR_INDEX.BLUE,
            "blue": WD_COLOR_INDEX.BLUE,
            "plum": WD_COLOR_INDEX.PINK,
            "pink": WD_COLOR_INDEX.PINK,
            "violet": WD_COLOR_INDEX.VIOLET,
            "mediumpurple": WD_COLOR_INDEX.VIOLET,
            "silver": WD_COLOR_INDEX.GRAY_25,
            "gray": WD_COLOR_INDEX.GRAY_50,
            "black": WD_COLOR_INDEX.BLACK,
            "navajowhite": WD_COLOR_INDEX.DARK_YELLOW,
            "rosybrown": WD_COLOR_INDEX.DARK_RED,
        }
        return mapping.get(name, WD_COLOR_INDEX.YELLOW)

    @staticmethod
    def _map_to_rgb(color_name: str | None, RGBColor) -> Optional[Any]:
        """
        Map a CSS-like color name to an RGBColor for text foreground (underline mode).

        Returns:
        - RGBColor or None if unknown or color_name is None.
        """
        if color_name is None:
            return None
        name = color_name.lower()
        hexmap = {
            "lightgreen": "90EE90",
            "orangered": "FF4500",
            "orange": "FFA500",
            "plum": "DDA0DD",
            "palegreen": "98FB98",
            "mediumseagreen": "3CB371",
            "steelblue": "4682B4",
            "skyblue": "87CEEB",
            "navajowhite": "FFDEAD",
            "mediumpurple": "9370DB",
            "rosybrown": "BC8F8F",
            "silver": "C0C0C0",
            "gray": "808080",
            "paleturquoise": "AFEEEE",
            "blue": "0000FF",
            "red": "FF0000",
            "green": "008000",
            "yellow": "FFFF00",
            "violet": "EE82EE",
            "pink": "FFC0CB",
            "turquoise": "40E0D0",
            "black": "000000",
        }
        hx = hexmap.get(name)
        if hx is None:
            return None
        r = int(hx[0:2], 16)
        g = int(hx[2:4], 16)
        b = int(hx[4:6], 16)
        return RGBColor(r, g, b)

# ---------- heatmap visualizer ----------
